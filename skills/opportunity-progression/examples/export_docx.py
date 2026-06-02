#!/usr/bin/env python3
"""
Export the Phase 7 report data to a Word document (.docx), styled to match the
Material Design 3 HTML/PDF output as closely as Word's model allows.

Design strategy:
    - Every card → a 1-cell borderless table with MD3 container background
    - Status chips → colored cells in inline tables
    - Section titles → MD3 primary purple with a thin purple underline
    - Zebra striping on data tables for readability
    - Scale focal numbers (score, stage completion) to 20-24pt for emphasis
    - Header section uses an elevated purple block so branding reads at a glance

Data source:
    Reads the same JSON data contract that feeds the Jinja2 template. No HTML
    parsing — single source of truth.

Usage:
    python3 examples/export_docx.py examples/sample-data.json
    # → writes sample-data.docx

Requires: pip install python-docx
"""
import json
import re
import sys
from pathlib import Path

# --- MD3 palette (mirrors :root in the Jinja2 template) ---------------------
MD3 = {
    "primary":                       "6750A4",
    "on_primary":                    "FFFFFF",
    "primary_container":             "EADDFF",
    "on_primary_container":          "21005D",
    "secondary_container":           "E8DEF8",
    "surface":                       "FFFBFE",
    "surface_container_low":         "F7F2FA",
    "surface_container":             "F3EDF7",
    "surface_container_high":        "ECE6F0",
    "surface_container_highest":     "E6E0E9",
    "on_surface":                    "1C1B1F",
    "on_surface_variant":            "49454F",
    "outline_variant":               "CAC4D0",
    "error":                         "B3261E",
    "error_container":               "F9DEDC",
    "on_error_container":            "410E0B",
    "warning":                       "8C5000",
    "warning_container":             "FFDDB1",
    "on_warning_container":          "2D1600",
    "success":                       "386A20",
    "success_container":             "B7F397",
    "on_success_container":          "0F2000",
}


def _set_normal_style_fonts(doc):
    """Set Normal style fonts via raw w:rFonts so Latin and CJK characters
    pick the correct typeface independently:
        ascii / hAnsi  → Amazon Ember (Latin)
        eastAsia / cs  → Source Han Sans SC (CJK)

    Word automatically routes characters to the right slot based on Unicode
    code point, so a single mixed-language paragraph will render Latin in
    Ember and CJK in Source Han Sans SC without manual run splitting.
    Missing fonts fall back to the OS default silently.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    style = doc.styles["Normal"]
    rpr = style.element.get_or_add_rPr()
    # remove any existing rFonts child to avoid duplicates
    existing = rpr.find(qn("w:rFonts"))
    if existing is not None:
        rpr.remove(existing)

    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Amazon Ember")
    rfonts.set(qn("w:hAnsi"), "Amazon Ember")
    rfonts.set(qn("w:eastAsia"), "Source Han Sans SC")
    rfonts.set(qn("w:cs"), "Source Han Sans SC")
    rpr.append(rfonts)


# --- python-docx helpers ----------------------------------------------------
def _shade(cell, fill_hex: str):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _borderless(cell):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "nil")
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def _cell_margins(cell, top=120, bottom=120, left=140, right=140):
    """Set cell padding in twentieths of a point (100 = 5pt)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for name, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tc_pr.append(tcMar)


def _run_color(run, hex_color: str):
    from docx.shared import RGBColor
    run.font.color.rgb = RGBColor.from_string(hex_color)


def _strip_html(s: str) -> str:
    return re.sub(r"<[^>]+>", "", s or "")


def _set_paragraph_bg(paragraph, fill_hex: str):
    """Shade an entire paragraph (used for the purple section dividers)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def _set_cell_vertical_align(cell, align="center"):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), align)
    tc_pr.append(vAlign)


def _add_bottom_border(paragraph, color_hex: str, size_pt: int = 12):
    """Add a colored bottom border to a paragraph — used as section divider."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size_pt))  # in eighths of a point
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


# --- Heading / Section helpers ---------------------------------------------
def add_section_title(doc, text):
    """Phase-level heading with a purple underline divider."""
    from docx.shared import Pt
    h = doc.add_paragraph()
    run = h.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    _run_color(run, MD3["primary"])
    _add_bottom_border(h, MD3["primary"], size_pt=12)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    return h


def add_sub_title(doc, text, size=13):
    from docx.shared import Pt
    h = doc.add_paragraph()
    run = h.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    _run_color(run, MD3["primary"])
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(2)
    return h


def add_card(doc, bg_hex=None):
    """Single-cell borderless table — our MD3 'card' surrogate."""
    from docx.shared import Pt
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    _borderless(cell)
    _cell_margins(cell, top=140, bottom=140, left=180, right=180)
    if bg_hex:
        _shade(cell, bg_hex)
    # remove the default empty paragraph so the first added content is the first line
    cell._tc.remove(cell.paragraphs[0]._p)
    return cell


def cell_add_run(cell, text, bold=False, italic=False, color=None, size=None):
    """Append a run to the last paragraph of a cell (or a new one if none)."""
    from docx.shared import Pt
    if not cell.paragraphs:
        cell.add_paragraph()
    p = cell.paragraphs[-1]
    r = p.add_run(text)
    if bold:
        r.bold = True
    if italic:
        r.italic = True
    if color:
        _run_color(r, color)
    if size:
        r.font.size = Pt(size)
    return r


def cell_add_para(cell, text="", bold=False, italic=False, color=None, size=None, space_before=None):
    from docx.shared import Pt
    p = cell.add_paragraph()
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if text:
        r = p.add_run(text)
        if bold:
            r.bold = True
        if italic:
            r.italic = True
        if color:
            _run_color(r, color)
        if size:
            r.font.size = Pt(size)
    return p


def add_kv(cell, key, value, key_color=None):
    p = cell.add_paragraph()
    kr = p.add_run(f"{key}: ")
    kr.bold = True
    if key_color:
        _run_color(kr, key_color)
    p.add_run(str(value))
    return p


def _tag_run(paragraph, text, fg_hex):
    """Inline 'chip' — Word lacks rounded backgrounds on runs, so we use
    bracket + colored bold as the chip simulation."""
    r = paragraph.add_run(f"  [{text}]")
    r.bold = True
    _run_color(r, fg_hex)
    return r


# --- Metric/Snapshot row (generic reusable pattern) ------------------------
def add_snapshot_row(doc, items, highlight_index=None):
    """Items: list of dicts with keys: label, value, color (optional),
    big (optional, bool — renders value at 16pt)."""
    from docx.shared import Pt
    t = doc.add_table(rows=1, cols=len(items))
    t.autofit = True
    for i, item in enumerate(items):
        c = t.rows[0].cells[i]
        _shade(c, MD3["surface_container_highest"])
        _borderless(c)
        _cell_margins(c, top=140, bottom=140, left=160, right=160)
        c.paragraphs[0].text = ""
        lr = c.paragraphs[0].add_run(item["label"].upper())
        lr.bold = True
        lr.font.size = Pt(8)
        _run_color(lr, MD3["on_surface_variant"])
        p = c.add_paragraph()
        vr = p.add_run(str(item["value"]))
        vr.bold = True
        vr.font.size = Pt(14 if item.get("big") else 12)
        _run_color(vr, item.get("color", MD3["on_surface"]))
    return t


# --- Section builders -------------------------------------------------------
def build_header(doc, meta):
    """Purple-container header block — reads as a branded title card."""
    from docx.shared import Pt
    cell = add_card(doc, bg_hex=MD3["primary_container"])
    _cell_margins(cell, top=220, bottom=220, left=280, right=280)

    # eyebrow label
    p = cell.add_paragraph()
    r = p.add_run("OPPORTUNITY PROGRESSION REPORT")
    r.bold = True
    r.font.size = Pt(9)
    _run_color(r, MD3["primary"])

    # title
    title_p = cell.add_paragraph()
    title_r = title_p.add_run(meta.get("opportunity_name", ""))
    title_r.bold = True
    title_r.font.size = Pt(22)
    _run_color(title_r, MD3["on_primary_container"])

    # subtitle
    parts = [meta.get("customer"), meta.get("industry"),
             f"{meta.get('tier', '')} Opportunity", meta.get("framework", "MEDDPICC")]
    subtitle = " · ".join([s for s in parts if s])
    if subtitle:
        sp = cell.add_paragraph()
        sr = sp.add_run(subtitle)
        sr.font.size = Pt(10)
        _run_color(sr, MD3["on_primary_container"])

    # ARR chip on the right (approximated by a bold ARR line)
    if meta.get("arr"):
        arr_p = cell.add_paragraph()
        ar = arr_p.add_run(f"ARR · {meta['arr']}")
        ar.bold = True
        ar.font.size = Pt(12)
        _run_color(ar, MD3["primary"])


def build_rollback(doc, rb):
    if not rb or not rb.get("show"):
        return
    from docx.shared import Pt
    cell = add_card(doc, bg_hex=MD3["warning_container"])
    p = cell.add_paragraph()
    r = p.add_run("⚠  ROLLBACK RECOMMENDED")
    r.bold = True
    r.font.size = Pt(12)
    _run_color(r, MD3["warning"])
    cell_add_para(cell,
                  f"分析继续在 {rb.get('current_stage', '')} 阶段展开。回退目标:{rb.get('target_stage', '')}",
                  bold=True, color=MD3["on_warning_container"], size=11)
    cell_add_para(cell, rb.get("reason", ""), color=MD3["on_warning_container"], size=10)


def build_phase2(doc, p2):
    if not p2:
        return
    add_section_title(doc, "Phase 2 · Deal Assessment")

    # verdict snapshot
    verdict_value = p2.get("verdict") or p2.get("health", "")
    verdict_color = MD3.get(p2.get("verdict_color") or p2.get("health_color", "primary"),
                            MD3["primary"])
    add_snapshot_row(doc, [
        {"label": "Verdict", "value": verdict_value, "color": verdict_color, "big": True},
        {"label": "Total Score", "value": p2.get("total", ""), "big": True},
        {"label": "Current Stage", "value": p2.get("stage", ""), "color": MD3["primary"], "big": True},
        {"label": "Classification", "value": p2.get("classification", ""), "big": True},
    ])

    # primary blocker — error container card
    if p2.get("primary_blocker"):
        pb = p2["primary_blocker"]
        cell = add_card(doc, bg_hex=MD3["error_container"])
        p = cell.add_paragraph()
        r = p.add_run("⛔  关键阻塞 · PRIMARY BLOCKER")
        r.bold = True
        _run_color(r, MD3["error"])
        cell_add_para(cell, pb.get("element", ""), bold=True, size=13,
                      color=MD3["on_error_container"])
        cell_add_para(cell, pb.get("rationale", ""), color=MD3["on_error_container"], size=10.5)

    # progression gaps — warning container card (only when present)
    if p2.get("progression_gaps"):
        pg = p2["progression_gaps"]
        cell = add_card(doc, bg_hex=MD3["warning_container"])
        p = cell.add_paragraph()
        r = p.add_run(f"📈  进下一阶段（{pg.get('target_stage', '')}）还需要")
        r.bold = True
        _run_color(r, MD3["warning"])
        for item in pg.get("items", []):
            bp = cell.add_paragraph(style="List Bullet")
            br = bp.add_run(item)
            from docx.shared import Pt
            br.font.size = Pt(10.5)
            _run_color(br, MD3["on_warning_container"])

    # Immediate Action + 30-Day Targets — side-by-side cards
    if p2.get("immediate_action") or p2.get("targets"):
        from docx.shared import Inches, Pt
        t = doc.add_table(rows=1, cols=2)
        t.autofit = False
        t.columns[0].width = Inches(3.1)
        t.columns[1].width = Inches(3.1)

        # immediate action
        c1 = t.rows[0].cells[0]
        _borderless(c1)
        _shade(c1, MD3["surface_container_highest"])
        _cell_margins(c1, top=160, bottom=160, left=180, right=180)
        c1.paragraphs[0].text = ""
        h = c1.paragraphs[0].add_run("🚀  本周行动 · IMMEDIATE ACTION")
        h.bold = True
        h.font.size = Pt(9)
        _run_color(h, MD3["primary"])
        if p2.get("immediate_action"):
            cell_add_para(c1, p2["immediate_action"], color=MD3["on_surface"], size=10.5)

        # 30-day targets
        c2 = t.rows[0].cells[1]
        _borderless(c2)
        _shade(c2, MD3["surface_container_highest"])
        _cell_margins(c2, top=160, bottom=160, left=180, right=180)
        c2.paragraphs[0].text = ""
        h = c2.paragraphs[0].add_run("🎯  30 天目标 · 30-DAY TARGETS")
        h.bold = True
        h.font.size = Pt(9)
        _run_color(h, MD3["primary"])
        for t_ in p2.get("targets", []):
            bp = c2.add_paragraph(style="List Bullet")
            run = bp.add_run(f"{t_['element']} ≥ {t_['threshold']}")
            run.bold = True
            _run_color(run, MD3["primary"])
            bp.add_run(f"  ·  {t_['text']}")

    # rationale — surface_container card (only when present)
    if p2.get("rationale"):
        cell = add_card(doc, bg_hex=MD3["surface_container"])
        cell_add_para(cell, "为什么这个判断", bold=True, color=MD3["on_surface_variant"], size=9)
        for para in p2["rationale"]:
            cell_add_para(cell, _strip_html(para), color=MD3["on_surface"], size=10.5)


def _zebra_table(doc, headers, rows):
    """Data table with purple header row and alternate-row striping.
    rows: list of list of (text, color_hex_or_None, bg_hex_or_None)."""
    from docx.shared import Pt
    table = doc.add_table(rows=1, cols=len(headers))

    # header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        _shade(hdr[i], MD3["primary"])
        _cell_margins(hdr[i], top=100, bottom=100, left=120, right=120)
        hdr[i].paragraphs[0].text = ""
        r = hdr[i].paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        _run_color(r, MD3["on_primary"])

    for ridx, row_cells in enumerate(rows):
        row = table.add_row().cells
        zebra_bg = MD3["surface_container_low"] if ridx % 2 == 0 else "FFFFFF"
        for ci, spec in enumerate(row_cells):
            if isinstance(spec, tuple):
                text, fg, bg_override = (spec + (None, None))[:3]
            else:
                text, fg, bg_override = spec, None, None
            _cell_margins(row[ci], top=80, bottom=80, left=120, right=120)
            bg = bg_override if bg_override else zebra_bg
            _shade(row[ci], bg)
            row[ci].paragraphs[0].text = ""
            rn = row[ci].paragraphs[0].add_run(str(text))
            rn.font.size = Pt(10)
            if fg:
                _run_color(rn, fg)
                rn.bold = True
    return table


def build_phase3(doc, p3):
    if not p3:
        return
    add_section_title(doc, "Phase 3 · Exit Criteria & Sales Stage Assessment")

    if p3.get("stage_aggregates"):
        from docx.shared import Pt
        t = doc.add_table(rows=1, cols=len(p3["stage_aggregates"]))
        t.autofit = True
        for i, sa in enumerate(p3["stage_aggregates"]):
            c = t.rows[0].cells[i]
            _shade(c, MD3["surface_container_highest"])
            _borderless(c)
            _cell_margins(c, top=160, bottom=160, left=180, right=180)
            c.paragraphs[0].text = ""
            # stage name + current tag
            np = c.paragraphs[0].add_run(sa["name"])
            np.bold = True
            np.font.size = Pt(11)
            _run_color(np, MD3["on_surface"])
            if sa.get("is_current"):
                tag = c.paragraphs[0].add_run("  · current")
                tag.italic = True
                tag.font.size = Pt(9)
                _run_color(tag, MD3["primary"])
            # big pct
            pp = c.add_paragraph()
            pct_run = pp.add_run(f"{sa['pct']}%")
            pct_run.bold = True
            pct_run.font.size = Pt(20)
            _run_color(pct_run, MD3.get(sa.get("color", "primary"), MD3["primary"]))
            # sublabel
            lp = c.add_paragraph()
            lr = lp.add_run(f"{sa['count']} exit criteria · 聚合完成度 {sa['pct']}%")
            lr.font.size = Pt(9)
            _run_color(lr, MD3["on_surface_variant"])

    # Exit criteria zebra table
    if p3.get("rows"):
        doc.add_paragraph()
        rows = []
        for r in p3["rows"]:
            status_color = r.get("status_color", "primary")
            bg = MD3.get(f"{status_color}_container", MD3["surface_container"])
            fg = MD3.get(f"on_{status_color}_container", MD3["on_surface"])
            rows.append([
                (str(r.get("num", "")),),
                (r.get("stage", ""),),
                (r.get("criterion", ""),),
                (f"{r.get('completion_pct', '')}%",),
                (f"{r.get('element_label', '')} ({r.get('element_pct', '')}%)",),
                (f"{r.get('status_icon', '')} {r.get('status', '')}", fg, bg),
            ])
        _zebra_table(doc, ["#", "Stage", "Exit Criterion", "Completion", "Mapped Element", "Status"], rows)

    # Risk cards row
    if p3.get("risks"):
        doc.add_paragraph()
        from docx.shared import Pt
        risks = p3["risks"]
        t = doc.add_table(rows=1, cols=len(risks))
        t.autofit = True
        for i, rk in enumerate(risks):
            c = t.rows[0].cells[i]
            bg_color = rk.get("bg_color", "error")
            _shade(c, MD3.get(f"{bg_color}_container", MD3["surface_container"]))
            _borderless(c)
            _cell_margins(c, top=140, bottom=140, left=160, right=160)
            c.paragraphs[0].text = ""
            tag = c.paragraphs[0].add_run(rk["kind"].upper())
            tag.bold = True
            tag.font.size = Pt(9)
            _run_color(tag, MD3.get(bg_color, MD3["error"]))
            p = c.add_paragraph()
            pr = p.add_run(rk.get("text", ""))
            pr.font.size = Pt(10)
            _run_color(pr, MD3.get(f"on_{bg_color}_container", MD3["on_surface"]))


def build_phase4(doc, p4):
    if not p4 or not p4.get("show"):
        return
    add_section_title(doc, "Phase 4 · Market & Competitive Intelligence")

    # 3-section grid
    from docx.shared import Pt
    t = doc.add_table(rows=1, cols=3)
    t.autofit = True
    sections = [
        ("1 · 客户当前云与 GenAI 栈", [item.get("text", "") for item in p4.get("current_stack", [])]),
        ("2 · 行业采购行为", None),  # handled specially
        ("3 · 同行业参考客户", ["★ " + ref for ref in p4.get("references", [])]),
    ]
    for i, (title, items) in enumerate(sections):
        c = t.rows[0].cells[i]
        _shade(c, MD3["surface_container_highest"])
        _borderless(c)
        _cell_margins(c, top=140, bottom=140, left=160, right=160)
        c.paragraphs[0].text = ""
        h = c.paragraphs[0].add_run(title)
        h.bold = True
        h.font.size = Pt(10)
        _run_color(h, MD3["primary"])
        if i == 1:
            bb = p4.get("buying_behavior", {})
            for k, label in [("factors", "决策因素"), ("painpoints", "采购痛点"),
                             ("cycle", "采购周期"), ("tech", "技术考虑")]:
                if bb.get(k):
                    p = c.add_paragraph()
                    lr = p.add_run(f"{label}: ")
                    lr.bold = True
                    lr.font.size = Pt(9.5)
                    v = p.add_run(bb[k])
                    v.font.size = Pt(9.5)
        else:
            for it in items or []:
                bp = c.add_paragraph()
                br = bp.add_run("• " + it)
                br.font.size = Pt(9.5)

    # Competitor cards
    if p4.get("competitors"):
        doc.add_paragraph()
        add_sub_title(doc, "Competitive Landscape")
        for c in p4["competitors"]:
            cell = add_card(doc, bg_hex=MD3["surface_container_high"])
            tp = cell.add_paragraph()
            nr = tp.add_run(c["name"])
            nr.bold = True
            nr.font.size = Pt(12)
            _run_color(nr, MD3["on_surface"])
            _tag_run(tp, c["threat_level"],
                     MD3.get(c.get("threat_color", "warning"), MD3["warning"]))
            cell_add_para(cell, c.get("desc", ""), color=MD3["on_surface"], size=10.5)
            add_kv(cell, "AWS 差异", c.get("aws_diff", ""), key_color=MD3["primary"])


def build_phase5(doc, p5):
    if not p5:
        return
    add_section_title(doc, "Phase 5 · Element Gap Analysis")

    # Layer A grid
    add_sub_title(doc, "Layer A · Element Summary")
    from docx.shared import Pt
    els = p5.get("elements", [])
    cols = 4
    for i in range(0, len(els), cols):
        row_els = els[i:i + cols]
        t = doc.add_table(rows=1, cols=cols)
        t.autofit = True
        for j in range(cols):
            c = t.rows[0].cells[j]
            _borderless(c)
            _cell_margins(c, top=140, bottom=140, left=140, right=140)
            c.paragraphs[0].text = ""
            if j < len(row_els):
                el = row_els[j]
                _shade(c, MD3["surface_container"])
                # code + priority tag
                p = c.paragraphs[0]
                cr = p.add_run(el["code"])
                cr.bold = True
                cr.font.size = Pt(14)
                _run_color(cr, MD3["on_surface"])
                pri_color = MD3.get(el.get("priority_color", "primary"), MD3["primary"])
                tag = p.add_run(f"    {el.get('priority', '')}")
                tag.bold = True
                tag.font.size = Pt(8)
                _run_color(tag, pri_color)
                # name
                np = c.add_paragraph()
                nr = np.add_run(el.get("name", ""))
                nr.font.size = Pt(9.5)
                _run_color(nr, MD3["on_surface_variant"])
                # score
                scr_color = MD3.get(el.get("score_color", "primary"), MD3["primary"])
                sp = c.add_paragraph()
                sr = sp.add_run(f"{el['score_pct']}%")
                sr.bold = True
                sr.font.size = Pt(18)
                _run_color(sr, scr_color)
                # relationship note
                cn = c.add_paragraph()
                cnr = cn.add_run(el.get("relationship_note", ""))
                cnr.font.size = Pt(8.5)
                _run_color(cnr, MD3["on_surface_variant"])

    # Layer B cards
    layer_b = [el for el in els if el.get("layer_b_card")]
    if layer_b:
        doc.add_paragraph()
        add_sub_title(doc, "Layer B · Diagnostic Cards")
        for el in layer_b:
            pri_color = MD3.get(el.get("priority_color", "primary"), MD3["primary"])
            cell = add_card(doc, bg_hex=MD3["surface_container"])
            p = cell.add_paragraph()
            hr = p.add_run(f"{el['code']} — {el['name']}")
            hr.bold = True
            hr.font.size = Pt(12)
            _run_color(hr, pri_color)
            sc = p.add_run(f"   ·   {el['score_pct']}%")
            sc.bold = True
            sc.font.size = Pt(12)
            _run_color(sc, MD3["on_surface_variant"])
            card = el["layer_b_card"]
            if card.get("evidence"):
                cell_add_para(cell, "Diagnostic Evidence: " + card["evidence"],
                              italic=True, color=MD3["on_surface_variant"], size=10)
            if card.get("root_cause"):
                add_kv(cell, "Root Cause", card["root_cause"], key_color=MD3["primary"])

    # Stakeholders — 2-col grid
    if p5.get("stakeholders"):
        doc.add_paragraph()
        add_sub_title(doc, "Stakeholder Profiling")
        if p5.get("stakeholder_source"):
            sp = doc.add_paragraph()
            sr = sp.add_run(p5["stakeholder_source"])
            sr.italic = True
            sr.font.size = Pt(9)
            _run_color(sr, MD3["on_surface_variant"])
        shs = p5["stakeholders"]
        for i in range(0, len(shs), 2):
            pair = shs[i:i + 2]
            t = doc.add_table(rows=1, cols=2)
            t.autofit = True
            for j in range(2):
                c = t.rows[0].cells[j]
                _borderless(c)
                _cell_margins(c, top=140, bottom=140, left=160, right=160)
                c.paragraphs[0].text = ""
                if j < len(pair):
                    sh = pair[j]
                    _shade(c, MD3["surface_container_highest"])
                    p = c.paragraphs[0]
                    nr = p.add_run(sh["name"])
                    nr.bold = True
                    nr.font.size = Pt(12)
                    _run_color(nr, MD3["on_surface"])
                    rl = p.add_run(f"    ·  {sh['role']}")
                    rl.font.size = Pt(10)
                    _run_color(rl, MD3["on_surface_variant"])
                    # tag as a second line to keep alignment clean
                    tag_fg_name = sh.get("tag_fg", "on_surface_variant").replace("-", "_")
                    tag_fg = MD3.get(tag_fg_name, MD3["on_surface_variant"])
                    tp = c.add_paragraph()
                    tr = tp.add_run(f"[ {sh['tag']} ]")
                    tr.bold = True
                    tr.font.size = Pt(9)
                    _run_color(tr, tag_fg)
                    dp = c.add_paragraph()
                    dr = dp.add_run(sh.get("desc", ""))
                    dr.font.size = Pt(10)
                    _run_color(dr, MD3["on_surface"])

    if p5.get("cxo_overlay"):
        cell = add_card(doc, bg_hex=MD3["primary_container"])
        p = cell.add_paragraph()
        h = p.add_run(f"👤  CxO Persona Overlay · {p5['cxo_overlay'].get('title', '')}")
        h.bold = True
        h.font.size = Pt(11)
        _run_color(h, MD3["on_primary_container"])
        cell_add_para(cell, p5["cxo_overlay"].get("text", ""),
                      color=MD3["on_primary_container"], size=10.5)

    # questions
    q = p5.get("questions")
    if q:
        doc.add_paragraph()
        add_sub_title(doc, "Customer Verification Questions")
        if q.get("persona_traits"):
            cell = add_card(doc, bg_hex=MD3["surface_container"])
            cell_add_para(cell, "PERSONA TRAITS APPLIED",
                          bold=True, color=MD3["on_surface_variant"], size=9)
            for idx, t in enumerate(q["persona_traits"], 1):
                bp = cell.add_paragraph()
                br = bp.add_run(f"{idx}. {t}")
                br.font.size = Pt(10)
        for block in q.get("by_element", []):
            pri_color = MD3.get(block.get("priority_color", "primary"), MD3["primary"])
            h = doc.add_paragraph()
            hr = h.add_run(block["element"])
            hr.bold = True
            hr.font.size = Pt(11)
            _run_color(hr, MD3["on_surface"])
            tag = h.add_run(f"   {block['priority']}")
            tag.bold = True
            tag.font.size = Pt(9)
            _run_color(tag, pri_color)
            h.paragraph_format.space_before = Pt(6)
            for idx, item in enumerate(block.get("items", []), 1):
                cell = add_card(doc, bg_hex=MD3["surface_container"])
                p = cell.add_paragraph()
                label = item.get("label") or f"主问题 {idx}"
                lr = p.add_run(f"{label}: ")
                lr.bold = True
                lr.font.size = Pt(10.5)
                _run_color(lr, MD3["primary"])
                qr = p.add_run(f'"{item.get("prompt", "")}"')
                qr.font.size = Pt(10.5)
                if item.get("purpose"):
                    add_kv(cell, "Purpose", item["purpose"])
                if item.get("insights"):
                    add_kv(cell, "Expected insights", item["insights"])
                if item.get("followup"):
                    add_kv(cell, "Follow-up", f'"{item["followup"]}"')


def build_phase6(doc, p6):
    if not p6:
        return
    from docx.shared import Pt
    add_section_title(doc, "Phase 6 · Action Plan")

    stage_val = p6.get("stage", "")
    if p6.get("stage_note"):
        stage_val += f" · {p6['stage_note']}"
    add_snapshot_row(doc, [
        {"label": "Score", "value": p6.get("score", ""), "color": MD3["on_surface"], "big": True},
        {"label": "Current Stage", "value": stage_val, "color": MD3["primary"], "big": True},
        {"label": "Completion", "value": f"{p6.get('completion_pct', '')}%",
         "color": MD3["warning"], "big": True},
    ])

    if p6.get("context"):
        cell = add_card(doc, bg_hex=MD3["surface_container"])
        cell_add_para(cell, "CONTEXT · 未完成出口标准",
                      bold=True, color=MD3["on_surface_variant"], size=9)
        for c in p6["context"]:
            bp = cell.add_paragraph()
            br = bp.add_run("• " + _strip_html(c))
            br.font.size = Pt(10.5)

    # Section 2 — Winning Strategies (one-liner overview, no nested weeks)
    if p6.get("strategies"):
        for idx, s in enumerate(p6.get("strategies", []), start=1):
            bg = MD3["primary_container"] if idx == 1 else MD3["surface_container_high"]
            fg = MD3["on_primary_container"] if idx == 1 else MD3["on_surface"]
            cell = add_card(doc, bg_hex=bg)
            _cell_margins(cell, top=140, bottom=140, left=200, right=200)
            tp = cell.add_paragraph()
            sid = s.get("id") or f"S{idx}"
            bd = tp.add_run(f"  {sid}  ")
            bd.bold = True
            bd.font.size = Pt(10)
            _run_color(bd, MD3["primary"])
            tr = tp.add_run(f"   {s.get('title', '')}")
            tr.bold = True
            tr.font.size = Pt(12)
            _run_color(tr, fg)
            if s.get("target"):
                tt = tp.add_run(f"   ·   {s['target']}")
                tt.bold = True
                tt.font.size = Pt(11)
                _run_color(tt, MD3["primary"])
            if s.get("tag"):
                _tag_run(tp, s["tag"], MD3["primary"])
            if s.get("rationale"):
                rp = cell.add_paragraph()
                rr = rp.add_run(s['rationale'])
                rr.italic = True
                rr.font.size = Pt(10)
                _run_color(rr, MD3["on_surface_variant"] if idx > 1 else MD3["on_primary_container"])

    # Section 3 — Weekly Plan (action-first, calendar-event de-duplicated)
    if p6.get("weekly_actions"):
        doc.add_paragraph()
        add_sub_title(doc, "Weekly Plan · 执行视角")
        for w in p6["weekly_actions"]:
            cell = add_card(doc, bg_hex=MD3["surface_container"])
            _cell_margins(cell, top=140, bottom=140, left=200, right=200)
            wh = cell.add_paragraph()
            wr = wh.add_run(w.get("week", ""))
            wr.bold = True
            wr.font.size = Pt(11)
            _run_color(wr, MD3["primary"])
            if w.get("label"):
                lr = wh.add_run(f"   ·   {w['label']}")
                lr.italic = True
                lr.font.size = Pt(9.5)
                _run_color(lr, MD3["on_surface_variant"])
            for a in w.get("actions", []):
                ap = cell.add_paragraph()
                ar = ap.add_run("→ " + a.get("description", ""))
                ar.font.size = Pt(10.5)
                _run_color(ar, MD3["on_surface"])
                tags = a.get("strategy_tags") or []
                if tags:
                    tg = ap.add_run(f"   [{' · '.join(tags)}]")
                    tg.bold = True
                    tg.font.size = Pt(9.5)
                    _run_color(tg, MD3["primary"])

        if p6.get("horizon_extension_reason"):
            hp = doc.add_paragraph()
            hr = hp.add_run(
                f"Horizon extended to {p6.get('weekly_horizon', 6)} weeks because "
                f"{p6['horizon_extension_reason']}"
            )
            hr.italic = True
            hr.font.size = Pt(9.5)
            _run_color(hr, MD3["on_surface_variant"])

    # Section 4 — Metrics (each as a small card)
    if p6.get("metrics"):
        doc.add_paragraph()
        add_sub_title(doc, "Key Success Metrics · 30-day targets")
        for m in p6["metrics"]:
            cell = add_card(doc, bg_hex=MD3["surface_container"])
            _cell_margins(cell, top=100, bottom=100, left=160, right=160)
            p = cell.add_paragraph()
            ck = p.add_run("✓  ")
            ck.bold = True
            ck.font.size = Pt(12)
            _run_color(ck, MD3["primary"])
            body = p.add_run(m["text"])
            body.font.size = Pt(10.5)
            tag_text = m.get("strategy_tag", "")
            if m.get("source_week"):
                tag_text += f"   ·   来源 {m['source_week']}"
            tag = p.add_run(f"   [{tag_text}]")
            tag.bold = True
            tag.font.size = Pt(9.5)
            _run_color(tag, MD3["primary"])


def build_footer(doc, meta):
    from docx.shared import Pt
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Generated by Opportunity Progression Skill · Dali Agent")
    r.italic = True
    r.font.size = Pt(9)
    _run_color(r, MD3["on_surface_variant"])
    if meta.get("report_date"):
        p2 = doc.add_paragraph()
        r2 = p2.add_run(f"Report Date: {meta['report_date']}")
        r2.font.size = Pt(8)
        _run_color(r2, MD3["on_surface_variant"])


# --- Top-level --------------------------------------------------------------
def build_document(data: dict, out_path: Path) -> Path:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        raise RuntimeError("python-docx is required. Install with: pip install python-docx")

    doc = Document()

    # tighter page margins so cards breathe naturally
    from docx.shared import Cm
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # base font — Latin: Amazon Ember; CJK: Source Han Sans SC
    # Set via w:rFonts so mixed-language runs route correctly per Unicode block.
    _set_normal_style_fonts(doc)
    style = doc.styles["Normal"]
    style.font.size = Pt(10.5)

    meta = data.get("meta", {})
    build_header(doc, meta)
    build_rollback(doc, data.get("rollback_banner"))
    build_phase2(doc, data.get("phase2"))
    build_phase3(doc, data.get("phase3"))
    build_phase4(doc, data.get("phase4"))
    build_phase5(doc, data.get("phase5"))
    build_phase6(doc, data.get("phase6"))
    build_footer(doc, meta)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python3 examples/export_docx.py <input.json> [output.docx]")
        sys.exit(2)

    src = Path(sys.argv[1]).resolve()
    if not src.exists():
        print(f"❌ Data JSON not found: {src}")
        sys.exit(1)

    out = Path(sys.argv[2]).resolve() if len(sys.argv) >= 3 else src.with_suffix(".docx")

    with open(src, "r", encoding="utf-8") as f:
        data = json.load(f)

    result = build_document(data, out)
    print(f"✅ DOCX exported: {result}")
